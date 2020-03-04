#-*- coding:utf-8 -*-
import os 
from collections import namedtuple
from urllib.request import urlopen
from urllib import parse
from bs4 import BeautifulSoup as bs
from docx import Document

__BASE_URL = "www.law.go.kr"
Pan = namedtuple("Pan", "pansi yozi jomun refpan num")

def __get_bs_obj(url):
    '''
    url (str)
    return beautifulSoup object
    '''
    url = "http://"+url
    html = urlopen(url)
    return bs(html, 'html.parser', from_encoding = 'utf-8')

def get_pan(pan_num):
    '''
    pan_num (str): id of precedent
    return Pan object
    '''
    url = "{}/판례/({})".format(__BASE_URL, pan_num)
    website = __get_bs_obj(parse.quote(url)) # handle the hangeul url
    pansrc = website.body.find('iframe', {'name':'lawService'})
    if pansrc is None:
        return Pan(None, None, None, None, pan_num)

    pansrc_url = pansrc['src']
    website = __get_bs_obj(__BASE_URL+"/"+pansrc_url)
    contents = website.body.find_all('p', {'class':'pty4'})
    contents = [c.text for c in contents]
    if len(contents) != 4:
        return Pan(None, None, None, None, pan_num)
    contents.append(pan_num)
    return Pan(*contents)

def save_pans(pans, ofilename, p_check=True, y_check=True, j_check=True):
    '''
    pans (list of Pan objects)
    ofilename (str) : file path to save pans as docx file
    '''
    doc = Document()
    for pan in pans:
        doc.add_heading(pan.num, level=1)
        if pan.pansi or pan.yozi or pan.jomun:
            if p_check and pan.pansi:
                doc.add_heading("판시사항", level=2)
                doc.add_paragraph(pan.pansi)
            if y_check and pan.yozi:
                doc.add_heading("판결요지", level=2)
                doc.add_paragraph(pan.yozi)
            if j_check and pan.jomun:
                doc.add_heading("참조조문", level=2)
                doc.add_paragraph(pan.jomun)
        else:
            doc.add_heading("판시사항, 판결요지, 참조조문 없습니다", level=2)
    doc.save(ofilename)
    print("save: {}".format(os.path.abspath(ofilename)))
